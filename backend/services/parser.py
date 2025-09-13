# Goal: Create a parser for DOCX and PDF files.
#
# Instructions:
# 1. Import the necessary libraries: `docx`, `pypdf`, and `io`.
# 2. Create functions to parse both DOCX and PDF files that take raw bytes.
# 3. For DOCX: Use `io.BytesIO(file_content)` to read the in-memory file with `docx.Document()`.
# 4. For PDF: Use `pypdf.PdfReader()` to read PDF content.
# 5. Create two dictionaries, one for `metadata` and one for `sections`.
# 6. Extract metadata and parse content by sections/headings.
# 7. Create a `full_text` key that contains all the text from the document joined together.
# 8. Return a single dictionary containing the metadata, sections, and full_text.

import docx
import pypdf
import io
import re

def parse_pdf(file_content: bytes) -> dict:
    """
    Parses a PDF file from bytes and extracts metadata, sections, and full text.
    
    Args:
        file_content (bytes): Raw bytes of the PDF file
    
    Returns:
        dict: Contains metadata, sections, and full_text
    """
    # Use BytesIO to read the in-memory file with pypdf.PdfReader()
    file_stream = io.BytesIO(file_content)
    pdf_reader = pypdf.PdfReader(file_stream)
    
    # Create dictionaries for metadata and sections
    metadata = {}
    sections = {}
    all_text = []
    
    # Extract metadata from PDF
    if pdf_reader.metadata:
        for key, value in pdf_reader.metadata.items():
            if value:
                clean_key = key.replace('/', '').strip()
                metadata[clean_key] = str(value).strip()
    
    # Add basic metadata
    metadata['total_pages'] = len(pdf_reader.pages)
    
    # Extract text from all pages
    full_page_text = []
    for page_num, page in enumerate(pdf_reader.pages):
        try:
            page_text = page.extract_text()
            if page_text.strip():
                full_page_text.append(page_text)
                all_text.append(page_text)
        except Exception as e:
            print(f"Error extracting text from page {page_num + 1}: {e}")
            continue
    
    # Parse sections based on common heading patterns
    full_text_content = '\n'.join(full_page_text)
    
    # Split text into lines for processing
    lines = full_text_content.split('\n')
    current_heading = "Introduction"  # Default section
    current_content = []
    
    # Look for heading patterns (all caps, numbered sections, etc.)
    heading_patterns = [
        r'^[A-Z][A-Z\s]{10,}$',  # All caps headings
        r'^\d+\.?\s+[A-Z]',      # Numbered sections (1. Introduction, 2. Methods)
        r'^[A-Z][a-z]+\s[A-Z]',  # Title Case headings
        r'^\d+\.\d+',            # Subsections (1.1, 2.1)
    ]
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line is a heading
        is_heading = False
        for pattern in heading_patterns:
            if re.match(pattern, line):
                is_heading = True
                break
        
        # Also check for short lines that might be headings (less than 50 chars, no periods at end)
        if not is_heading and len(line) < 50 and not line.endswith('.') and len(line.split()) <= 6:
            # Check if it's likely a heading (starts with capital, has capital letters)
            if line[0].isupper() and any(c.isupper() for c in line[1:]):
                is_heading = True
        
        if is_heading:
            # Save previous section if it has content
            if current_content:
                sections[current_heading] = '\n'.join(current_content)
            
            # Start new section
            current_heading = line
            current_content = []
        else:
            # Add content to current section
            current_content.append(line)
    
    # Don't forget the last section
    if current_content:
        sections[current_heading] = '\n'.join(current_content)
    
    # Create full_text by joining all text
    full_text = '\n'.join(all_text)
    
    # Return the complete parsed document
    return {
        'metadata': metadata,
        'sections': sections,
        'full_text': full_text
    }

def parse_docx(file_content: bytes) -> dict:
    """
    Parses a DOCX file from bytes and extracts metadata, sections, and full text.
    
    Args:
        file_content (bytes): Raw bytes of the DOCX file
    
    Returns:
        dict: Contains metadata, sections, and full_text
    """
    # Use BytesIO to read the in-memory file with docx.Document()
    file_stream = io.BytesIO(file_content)
    document = docx.Document(file_stream)
    
    # Create dictionaries for metadata and sections
    metadata = {}
    sections = {}
    all_text = []
    
    # Extract paragraphs
    paragraphs = document.paragraphs
    
    # Extract metadata from first few lines (first 5 paragraphs or until first heading)
    metadata_lines = []
    metadata_count = 0
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # Stop extracting metadata if we hit a heading or after 5 lines
        if para.style.name.startswith('Heading') or metadata_count >= 5:
            break
            
        # Look for common metadata patterns
        if any(keyword in text.lower() for keyword in ['document id', 'version', 'date', 'author', 'title']):
            # Try to extract key-value pairs
            if ':' in text:
                key, value = text.split(':', 1)
                metadata[key.strip()] = value.strip()
            else:
                metadata[f'metadata_{metadata_count}'] = text
        else:
            metadata[f'metadata_{metadata_count}'] = text
            
        metadata_count += 1
    
    # Parse the rest of the document by heading titles
    current_heading = "Introduction"  # Default section for content before first heading
    current_content = []
    
    for para in paragraphs:
        text = para.text.strip()
        
        # Skip empty paragraphs
        if not text:
            continue
            
        # Add to full text
        all_text.append(text)
        
        # Check if this paragraph is a heading
        if para.style.name.startswith('Heading'):
            # Save previous section if it has content
            if current_content:
                sections[current_heading] = '\n'.join(current_content)
            
            # Start new section
            current_heading = text
            current_content = []
        else:
            # Add content to current section (skip if it's already in metadata)
            if text not in metadata.values():
                current_content.append(text)
    
    # Don't forget the last section
    if current_content:
        sections[current_heading] = '\n'.join(current_content)
    
    # Create full_text by joining all text
    full_text = '\n'.join(all_text)
    
    # Return the complete parsed document
    return {
        'metadata': metadata,
        'sections': sections,
        'full_text': full_text
    }

def parse_document(file_content: bytes, filename: str = None) -> dict:
    """
    Universal parser that can handle both DOCX and PDF files.
    
    Args:
        file_content (bytes): Raw bytes of the file
        filename (str, optional): Original filename to help determine file type
    
    Returns:
        dict: Contains metadata, sections, and full_text
    
    Raises:
        ValueError: If file type is not supported or cannot be determined
    """
    # Determine file type from filename extension if provided
    file_type = None
    if filename:
        filename_lower = filename.lower()
        if filename_lower.endswith('.docx'):
            file_type = 'docx'
        elif filename_lower.endswith('.pdf'):
            file_type = 'pdf'
    
    # If we couldn't determine from filename, try to detect from content
    if not file_type:
        # Check PDF magic number
        if file_content.startswith(b'%PDF'):
            file_type = 'pdf'
        # Check for DOCX magic number (ZIP file with specific structure)
        elif file_content.startswith(b'PK'):
            # DOCX files are ZIP archives, so they start with PK
            # We'll assume it's DOCX if it starts with PK and we don't know otherwise
            file_type = 'docx'
    
    # Parse based on detected file type
    if file_type == 'docx':
        return parse_docx(file_content)
    elif file_type == 'pdf':
        return parse_pdf(file_content)
    else:
        raise ValueError("Unsupported file type. Only DOCX and PDF files are supported.")
